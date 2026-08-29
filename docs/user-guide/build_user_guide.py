from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SCREENSHOTS = ROOT / "docs" / "user-guide" / "screenshots"
OUTPUT = ROOT / "docs" / "Pulse_CRM_Руководство_пользователя.docx"

FONT = "Calibri"
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
ACCENT = "4F46E5"
ACCENT_SOFT = "F1F0FF"
INK = "172033"
MUTED = "667085"
LIGHT = "F4F6F9"
TABLE_FILL = "E8EEF5"
BORDER = "D9DEE8"
WARNING_FILL = "FFF7ED"
WARNING = "C2410C"
SUCCESS_FILL = "F0FDF4"
SUCCESS = "15803D"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_style_font(style, name: str, size: float | None = None, color: str | None = None,
                   bold: bool | None = None, italic: bool | None = None) -> None:
    style.font.name = name
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def set_run_font(run, *, name: str = FONT, size: float | None = None,
                 color: str | None = None, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr()
    rfonts = run._element.rPr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        run._element.rPr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, FONT, 11, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    set_style_font(title, FONT, 31, NAVY, bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, FONT, 14, DARK_BLUE)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.line_spacing = 1.15

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        set_style_font(style, FONT, size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.05
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = doc.styles["Caption"]
    set_style_font(caption, FONT, 9, MUTED, italic=True)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.keep_together = True

    custom = {
        "Guide Kicker": (10, ACCENT, True, False, WD_ALIGN_PARAGRAPH.CENTER, 0, 14, 1.0),
        "Guide Lead": (11.5, DARK_BLUE, False, False, WD_ALIGN_PARAGRAPH.LEFT, 0, 10, 1.25),
        "Guide Small": (9, MUTED, False, False, WD_ALIGN_PARAGRAPH.LEFT, 0, 4, 1.1),
        "Guide Step": (10.5, INK, False, False, WD_ALIGN_PARAGRAPH.LEFT, 0, 4, 1.2),
        "Guide Callout": (10.5, INK, False, False, WD_ALIGN_PARAGRAPH.LEFT, 4, 4, 1.2),
    }
    for name, (size, color, bold, italic, align, before, after, spacing) in custom.items():
        if name not in doc.styles:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = doc.styles[name]
        set_style_font(style, FONT, size, color, bold=bold, italic=italic)
        style.paragraph_format.alignment = align
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = spacing


def add_numbering_definition(doc: Document, fmt: str, text: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    lvl.append(ppr)
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT)
    rpr.append(rfonts)
    lvl.append(rpr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)


def add_list_item(doc: Document, text: str, num_id: int, *, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="Guide Step")
    apply_num(p, num_id)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    return p


def shade_paragraph(paragraph, fill: str, border_color: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "14")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_color)
    pbdr.append(left)
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(7)


def add_callout(doc: Document, label: str, text: str, *, kind: str = "info"):
    colors = {
        "info": (LIGHT, ACCENT),
        "warning": (WARNING_FILL, WARNING),
        "success": (SUCCESS_FILL, SUCCESS),
    }
    fill, border = colors[kind]
    p = doc.add_paragraph(style="Guide Callout")
    r = p.add_run(f"{label}: ")
    set_run_font(r, bold=True, color=border)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    shade_paragraph(p, fill, border)
    return p


def set_cell_margins(cell) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for side, value in CELL_MARGINS_DXA.items():
        node = tcmar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError("Table widths must total 9360 DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tblpr = tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tblind.set(qn("w:type"), "dxa")
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths_dxa[index]))
            tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_table_borders(table, color: str = BORDER) -> None:
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        borders.append(node)


def shade_cell(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, TABLE_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=NAVY, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(text)
            set_run_font(r, size=9.5, color=INK)
    set_table_geometry(table, widths_dxa)
    return table


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("PULSE CRM  |  РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fr = fp.add_run("Страница ")
    set_run_font(fr, size=8.5, color=MUTED)
    add_page_field(fp)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = ""


def add_spacer(doc: Document, points: float) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(points)
    p.paragraph_format.space_after = Pt(0)


def add_cover(doc: Document) -> None:
    add_spacer(doc, 58)
    p = doc.add_paragraph(style="Guide Kicker")
    p.add_run("РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ")
    p = doc.add_paragraph(style="Title")
    p.add_run("Pulse CRM")
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("Сделки, клиенты, задачи и администрирование в одном рабочем пространстве")
    add_spacer(doc, 30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("От первого обращения - до следующей покупки")
    set_run_font(r, size=15, color=ACCENT, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.65)
    p.paragraph_format.right_indent = Inches(0.65)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(
        "Пошаговое руководство для менеджеров, администраторов и владельцев: "
        "вход, ежедневная работа и основные настройки."
    )
    set_run_font(r, size=11.5, color=INK)
    add_spacer(doc, 52)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Версия MVP 0.1  |  28 августа 2026")
    set_run_font(r, size=10, color=MUTED, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Скриншоты подготовлены в демонстрационном рабочем пространстве; имена и суммы приведены для примера.")
    set_run_font(r, size=9, color=MUTED, italic=True)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_title_block(doc: Document, title: str, lead: str | None = None) -> None:
    p = doc.add_paragraph(title, style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    if lead:
        p = doc.add_paragraph(lead, style="Guide Lead")
        p.paragraph_format.keep_with_next = True


def add_screenshot(doc: Document, filename: str, caption: str, alt_text: str,
                   *, max_width: float = 6.35, max_height: float = 5.8) -> None:
    path = SCREENSHOTS / filename
    with Image.open(path) as image:
        width_px, height_px = image.size
    ratio = width_px / height_px
    width_in = min(max_width, max_height * ratio)
    height_in = width_in / ratio
    if height_in > max_height:
        height_in = max_height
        width_in = height_in * ratio
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width_in), height=Inches(height_in))
    shape._inline.docPr.set("name", caption)
    shape._inline.docPr.set("title", caption)
    shape._inline.docPr.set("descr", alt_text)
    cp = doc.add_paragraph(caption, style="Caption")
    cp.paragraph_format.keep_with_next = True


def add_intro_page(doc: Document, bullet_num: int) -> None:
    add_page_break(doc)
    add_title_block(
        doc,
        "Перед началом",
        "Это руководство описывает ежедневную работу в Pulse CRM и административные функции, доступные владельцу и администратору.",
    )
    doc.add_paragraph("Роли и доступ", style="Heading 2")
    add_table(
        doc,
        ["Роль", "Что видит", "Ключевые права"],
        [
            ["Владелец", "Все разделы", "CRM-данные, все настройки, приглашение менеджеров и администраторов"],
            ["Администратор", "Все разделы", "CRM-данные, воронки, поля, каналы, оповещения, импорт; приглашение менеджеров"],
            ["Менеджер", "Без «Настроек»", "Сделки, клиенты, задачи, переписка и активность"],
        ],
        [1650, 2050, 5660],
    )
    doc.add_paragraph("Как читать руководство", style="Heading 2")
    for text in (
        "Менеджеру: последовательно пройти разделы «Вход» - «Активность».",
        "Владельцу или администратору: дополнительно изучить раздел «Настройки».",
        "На телефоне: использовать нижнее меню; этап сделки менять через список в карточке.",
    ):
        add_list_item(doc, text, bullet_num)
    add_callout(
        doc,
        "Важно",
        "Интерфейс может незначительно отличаться в зависимости от роли, подключённых каналов и размера экрана.",
    )


def add_screen_page(doc: Document, title: str, lead: str, image: str, caption: str,
                    alt: str, items: list[str], num_id: int, *, ordered: bool = False,
                    callout: tuple[str, str, str] | None = None,
                    max_width: float = 6.35, max_height: float = 5.8) -> None:
    add_page_break(doc)
    add_title_block(doc, title, lead)
    add_screenshot(doc, image, caption, alt, max_width=max_width, max_height=max_height)
    list_id = add_numbering_definition(doc, "decimal", "%1.") if ordered else num_id
    for item in items:
        add_list_item(doc, item, list_id)
    if callout:
        label, text, kind = callout
        add_callout(doc, label, text, kind=kind)


def build_document() -> None:
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    bullet_num = add_numbering_definition(doc, "bullet", "•")
    decimal_num = add_numbering_definition(doc, "decimal", "%1.")

    doc.core_properties.title = "Pulse CRM - руководство пользователя"
    doc.core_properties.subject = "Пошаговая работа со сделками, клиентами, задачами и настройками"
    doc.core_properties.author = "Pulse CRM"
    doc.core_properties.keywords = "Pulse CRM, CRM, руководство пользователя"
    doc.core_properties.created = datetime(2026, 8, 28, tzinfo=timezone.utc)
    doc.core_properties.modified = datetime(2026, 8, 28, tzinfo=timezone.utc)

    add_cover(doc)
    add_intro_page(doc, bullet_num)

    add_screen_page(
        doc,
        "Вход в Pulse CRM",
        "Администратор выдаёт доступ или присылает ссылку-приглашение. После входа открывается рабочее пространство компании.",
        "00-login.jpg",
        "Рисунок 1. Форма входа",
        "Форма входа Pulse CRM с полями Email и Пароль",
        [
            "Откройте адрес Pulse CRM, который выдал администратор.",
            "Введите Email и пароль, затем нажмите «Войти».",
            "Если вы получили приглашение, перейдите по ссылке, укажите имя и задайте пароль не короче 12 символов.",
        ],
        decimal_num,
        ordered=True,
        callout=("Совет", "Не передавайте пароль коллегам. Каждому сотруднику нужна отдельная учётная запись.", "info"),
    )

    add_screen_page(
        doc,
        "Навигация и уведомления",
        "На компьютере основные разделы расположены слева. В правом верхнем углу находятся уведомления и меню аккаунта.",
        "01-notifications.jpg",
        "Рисунок 2. Центр уведомлений",
        "Главная страница Pulse CRM с открытым меню уведомлений",
        [
            "Меню слева: «Главная», «Сделки», «Клиенты», «Задачи», «Активность» и, для административных ролей, «Настройки».",
            "Колокольчик: показывает новые лиды, просроченные задачи и напоминания о следующей покупке.",
            "Аватар: открывает данные аккаунта и команду «Выйти».",
        ],
        bullet_num,
        max_width=5.85,
    )

    add_screen_page(
        doc,
        "Главная: что требует внимания",
        "Главная страница собирает просроченные задачи, ключевые метрики, свежие события и состояние воронки.",
        "01-home.jpg",
        "Рисунок 3. Главная страница",
        "Главная Pulse CRM с метриками, задачами и последней активностью",
        [
            "Начните с оранжевого блока: он показывает срочную проблему и ведёт к задачам.",
            "Метрики «Активные сделки», «Новые обращения» и «Следующие покупки» помогают расставить приоритеты.",
            "Ссылки «Все задачи», «Вся история» и «Открыть воронку» ведут к детальным спискам.",
        ],
        bullet_num,
        callout=("Рабочий путь", "Новую сделку в текущем MVP создавайте через «Сделки» - «Новая сделка».", "info"),
    )

    add_screen_page(
        doc,
        "Сделки: Kanban-воронка",
        "Воронка показывает, на каком этапе находится каждая сделка. На карточке видны сумма, источник, ответственный и срок.",
        "02-deals-kanban.jpg",
        "Рисунок 4. Kanban-воронка сделок",
        "Kanban-воронка Pulse CRM с четырьмя этапами и карточками сделок",
        [
            "Выберите воронку в списке слева от строки поиска.",
            "Найдите сделку по названию, потребности или источнику.",
            "На компьютере перетащите карточку в нужную колонку. На телефоне меняйте этап в карточке сделки.",
        ],
        decimal_num,
        ordered=True,
        callout=("Если переход не выполнился", "Заполните поля, которые CRM показала как обязательные, и повторите смену этапа.", "warning"),
        max_width=5.75,
    )

    add_screen_page(
        doc,
        "Создание сделки",
        "Новая сделка попадает на первый этап выбранной воронки.",
        "03-new-deal.jpg",
        "Рисунок 5. Форма «Новая сделка»",
        "Модальная форма создания сделки с полями названия, потребности, суммы и источника",
        [
            "Откройте «Сделки» и нажмите «Новая сделка».",
            "Заполните «Название», «Потребность» и «Сумма, ₽».",
            "Укажите источник: ручной ввод, Email, Telegram, MAX, Webhook или HTML-форма.",
            "Нажмите «Создать сделку» и откройте её карточку для дополнения.",
        ],
        decimal_num,
        ordered=True,
    )

    add_screen_page(
        doc,
        "Список, поиск и фильтры сделок",
        "Переключатель в верхней части экрана меняет Kanban на табличный список. Фильтр по источнику помогает быстро выделить нужные обращения.",
        "06-deals-list-filter.jpg",
        "Рисунок 6. Список сделок с фильтром Telegram",
        "Табличный список сделок Pulse CRM с открытм фильтром источника Telegram",
        [
            "Кнопка «Список» показывает этап, сумму, источник, срок и ответственного в одной строке.",
            "Нажмите «Фильтры» и выберите источник или оставьте «Все источники».",
            "Нажмите на строку, чтобы открыть карточку сделки.",
        ],
        bullet_num,
    )

    add_screen_page(
        doc,
        "Карточка сделки",
        "В карточке собраны контакты, этап, ответственный, следующая покупка, задачи, переписка и история.",
        "04-deal-details.jpg",
        "Рисунок 7. Карточка сделки на вкладке «Детали»",
        "Карточка сделки Кофейня Слой с деталями, задачей и перепиской",
        [
            "Выберите «Этап сделки», если её нужно переместить без перетаскивания.",
            "Укажите дату «Следующая покупка» и нажмите «Сохранить».",
            "Нажатие на задачу меняет её состояние между открытой и выполненной.",
        ],
        bullet_num,
        callout=("Важно", "Дата следующей покупки создаёт напоминание, но не создаёт новую сделку автоматически.", "warning"),
    )

    add_screen_page(
        doc,
        "Переписка, вложения и история",
        "Вкладка «Переписка» хранит входящие и исходящие сообщения по подключённому каналу.",
        "05-deal-messages.jpg",
        "Рисунок 8. Переписка в карточке сделки",
        "Вкладка переписки сделки с входящим и исходящим сообщением",
        [
            "Откройте «Переписка», напишите текст и нажмите «Отправить».",
            "Для вложения нажмите скрепку. Допустимы изображения, PDF, TXT/CSV и распространённые офисные форматы до 20 МБ.",
            "Если отправка завершилась ошибкой, используйте «Повторить» после устранения причины.",
            "Во вкладке «История» можно добавить заметку; записи в ленте не редактируются.",
        ],
        bullet_num,
        callout=("Если кнопка отправки недоступна", "Попросите администратора проверить статус канала в «Настройки» - «Каналы».", "warning"),
        max_width=5.45,
    )

    add_screen_page(
        doc,
        "Клиенты: контакты и компании",
        "Раздел «Клиенты» объединяет людей и организации. Переключатель справа меняет «Контакты» на «Компании».",
        "07-clients.jpg",
        "Рисунок 9. Список контактов",
        "Список клиентов Pulse CRM с контактами, покупками, следующей покупкой и ответственным",
        [
            "Используйте поиск по имени, компании, телефону или Email.",
            "Нажмите «Новый контакт» и укажите имя, фамилию, Email и телефон.",
            "Для компании переключите вид на «Компании», затем укажите название, Email, телефон и сайт.",
        ],
        bullet_num,
    )

    add_screen_page(
        doc,
        "Карточка контакта",
        "Карточка контакта показывает связи, накопленную выручку, число покупок, плановую дату и историю.",
        "08-contact-card.jpg",
        "Рисунок 10. Карточка контакта",
        "Карточка контакта Анны Смирновой с контактными данными и историей",
        [
            "Нажмите на строку клиента, чтобы открыть карточку.",
            "Блок «История покупок» собирает выигранные сделки, связанные с контактом.",
            "В поле «Добавить заметку» зафиксируйте договорённость и нажмите «Добавить».",
        ],
        decimal_num,
        ordered=True,
        max_width=5.65,
    )

    add_screen_page(
        doc,
        "Задачи",
        "Раздел «Задачи» помогает планировать звонки, встречи и напоминания о следующей покупке.",
        "09-new-task.jpg",
        "Рисунок 11. Форма «Новая задача»",
        "Форма новой задачи с типом, сроком, напоминанием, исполнителем и сделкой",
        [
            "Отфильтруйте список: «Все», «Сегодня», «Просрочено» или «Предстоящие»; для точного поиска используйте строку поиска.",
            "Нажмите «Новая задача» и заполните название, тип, срок, напоминание и исполнителя.",
            "При необходимости выберите сделку или оставьте «Без привязки», затем нажмите «Создать задачу».",
            "Нажмите на строку задачи, чтобы отметить её выполненной; повторное нажатие вернёт её в работу.",
        ],
        decimal_num,
        ordered=True,
    )

    add_screen_page(
        doc,
        "Активность",
        "Журнал активности - это неизменяемая история действий команды и интеграций.",
        "10-activity.jpg",
        "Рисунок 12. Журнал активности",
        "Журнал активности Pulse CRM с поиском, фильтром типа события и лентой",
        [
            "Найдите событие по тексту в строке «Поиск по событиям».",
            "Откройте «Фильтры» и выберите: все события, сделки, сообщения, задачи, контакты или система.",
            "Используйте журнал для проверки, кто и когда создал клиента, переместил сделку, выполнил задачу или получил сообщение.",
        ],
        bullet_num,
        max_width=5.45,
        max_height=5.55,
    )

    add_screen_page(
        doc,
        "Настройки: воронки и поля",
        "Раздел «Настройки» виден только владельцу и администратору. Изменения применяются ко всей компании.",
        "11-settings-pipelines.jpg",
        "Рисунок 13. Воронки и поля",
        "Настройки Pulse CRM с пользовательскими полями и этапами воронки",
        [
            "«Новая воронка» создаёт готовый набор открытых и системных этапов.",
            "Карандаш у воронки меняет её название; кнопка «+» добавляет новый открытый этап.",
            "У этапа можно изменить название. Удаляется только пустой открытый этап; системные этапы и последний открытый этап защищены.",
            "Корзина удаляет только не последнюю воронку без связанных сделок, каналов, форм, webhook-ов и правил оповещений.",
            "«Новое поле сделки» добавляет текст, число, дату, флаг или список в карточку сделки.",
            "Нажмите на этап, чтобы определить поля, без которых переход дальше будет заблокирован.",
        ],
        bullet_num,
        callout=(
            "Важно",
            "Входящий лид не теряется из-за незаполненных полей. Если удаление отклонено, перенесите сделки и перенастройте указанные зависимости, затем повторите действие.",
            "warning",
        ),
    )

    add_screen_page(
        doc,
        "Обязательные поля этапа",
        "Выберите встроенные и пользовательские поля, которые менеджер должен заполнить до перехода на следующий этап.",
        "12-settings-required-fields.jpg",
        "Рисунок 14. Редактор обязательных полей",
        "Форма настройки обязательных полей этапа Предложение",
        [
            "Отметьте нужные поля в группах «Встроенные поля» и «Пользовательские поля сделки».",
            "Нажмите «Сохранить обязательность».",
        ],
        decimal_num,
        ordered=True,
        max_width=5.30,
        max_height=5.65,
    )

    add_screen_page(
        doc,
        "Пользователи и роли",
        "Приглашение действует 72 часа. Владелец может приглашать менеджеров и администраторов; администратор - только менеджеров.",
        "13-settings-users.jpg",
        "Рисунок 15. Пользователи и приглашения",
        "Настройки пользователей Pulse CRM с формой приглашения сотрудника",
        [
            "Откройте «Настройки» - «Пользователи».",
            "Введите Email сотрудника и выберите роль.",
            "Нажмите «Создать приглашение», скопируйте готовую ссылку и передайте её только адресату.",
        ],
        decimal_num,
        ordered=True,
        callout=("Безопасность", "Ссылка даёт доступ к созданию учётной записи. Не публикуйте её в общих чатах или открытых документах.", "warning"),
    )

    add_screen_page(
        doc,
        "Каналы и источники",
        "Каналы связывают входящие обращения и исходящую переписку с воронкой, начальным этапом и, при необходимости, ответственным.",
        "14-settings-channels.jpg",
        "Рисунок 16. Общие каналы компании",
        "Список каналов Pulse CRM: Email, Telegram, MAX, Webhook и HTML-форма",
        [
            "«Подключить канал»: общая почта IMAP/SMTP, Telegram-бот или MAX-бот.",
            "«Webhook»: HMAC-защищённый endpoint для внешней системы.",
            "«HTML-форма»: форма на сайте с ограничением разрешённых доменов.",
            "Перед активацией проверьте реквизиты, маршрут в воронку и тестовое обращение.",
        ],
        bullet_num,
        callout=("Безопасность", "Токены, пароли и секреты вводите только в защищённых полях Pulse CRM; не добавляйте их в заметки или общие чаты.", "warning"),
        max_width=5.45,
    )

    add_screen_page(
        doc,
        "Правила оповещений",
        "Правило связывает событие, получателя, канал, задержку и текст шаблона.",
        "15-settings-notifications.jpg",
        "Рисунок 17. Правила оповещений",
        "Список включённых и выключенных правил оповещений Pulse CRM",
        [
            "Нажмите «Новое правило», укажите событие, получателя и канал: в приложении, Email, Telegram или MAX.",
            "Для оповещения клиента зафиксируйте основание и подтверждённое согласие на конкретный канал.",
            "После теста включите правило переключателем; таким же способом его можно временно выключить.",
        ],
        decimal_num,
        ordered=True,
        callout=("Важно", "Клиентское правило создавайте выключенным и активируйте только после проверки согласия и адреса.", "warning"),
    )

    add_screen_page(
        doc,
        "Импорт из amoCRM",
        "Импорт предназначен для разового контролируемого переноса воронок, контактов, компаний, сделок, задач, полей и заметок.",
        "16-settings-import.jpg",
        "Рисунок 18. Начальный экран импорта amoCRM",
        "Настройки импорта из amoCRM с кнопкой подключения и блоками статуса",
        [
            "Сначала нажмите «Подключить amoCRM» и заполните реквизиты OAuth-интеграции.",
            "Запустите dry-run без записи бизнес-данных и проверьте сопоставление воронок и пользователей.",
            "После проверки запустите перенос; следите за прогрессом и ошибками, при необходимости используйте паузу, продолжение и повтор.",
        ],
        decimal_num,
        ordered=True,
        callout=("Перед импортом", "Администратор должен создать резервную копию и выполнить перенос в запланированное окно.", "warning"),
        max_width=5.45,
    )

    add_page_break(doc)
    add_title_block(
        doc,
        "Мобильная работа и быстрая помощь",
        "Pulse CRM адаптируется к телефону и планшету. На маленьом экране основная навигация переносится вниз.",
    )
    doc.add_paragraph("На телефоне", style="Heading 2")
    for text in (
        "Используйте нижнее меню: «Главная», «Сделки», «Клиенты», «Задачи» и «Ещё».",
        "Открывайте сделку касанием; для смены этапа используйте список «Этап сделки», а не перетаскивание.",
        "Пункт «Ещё» ведёт в «Настройки» для владельца/администратора и в «Активность» для менеджера.",
    ):
        add_list_item(doc, text, bullet_num)

    doc.add_paragraph("Типовые ситуации", style="Heading 2")
    add_table(
        doc,
        ["Ситуация", "Что делать"],
        [
            ["Нет «Настроек»", "Это нормально для роли «Менеджер»; обратитесь к владельцу или администратору."],
            ["Сделка не меняет этап", "Заполните перечисленные CRM обязательные поля и повторите действие."],
            ["Данные устарели", "Обновите страницу, снова откройте карточку и повторите изменение."],
            ["Сообщение не отправляется", "Проверьте канал и адрес получателя; после исправления нажмите «Повторить»."],
            ["Вложение отклонено", "Проверьте формат и размер: архивы и исполняемые файлы не принимаются, лимит - 20 МБ."],
        ],
        [2850, 6510],
    )
    add_callout(
        doc,
        "Результат",
        "Для ежедневной работы достаточно цикла: «Главная» - срочные задачи - «Сделки» - обновление карточек - «Активность» для контроля.",
        kind="success",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
